import pandas as pd
from collections import Counter
import numpy as np
from docx import Document
import time
from docx.shared import Cm, Mm
from docx.enum.text import WD_LINE_SPACING

import docx
# from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from docx.shared import RGBColor
from docx.oxml.shared import OxmlElement


def set_title_font(pargraph):
    pargraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    runs = pargraph.runs

    for header_run in runs:
        # 设置西文字体：Times New Roman
        header_run.font.name = 'Times New Roman'
        # 设置中文字体：宋体
        header_run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 设置字体大小：小三equal pt15
        header_run.font.size = Pt(15)
        # 设置加粗：是
        header_run.font.bold = True
        # 设置字体颜色：黑色
        header_run.font.color.rgb = RGBColor(0, 0, 0)
        # 设置下划线： False
        header_run.font.underline = False


def set_heading_font(pargraph):
    runs = pargraph.runs

    for header_run in runs:
        # 设置西文字体：Times New Roman
        header_run.font.name = 'Times New Roman'
        # 设置中文字体：宋体
        header_run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 设置字体大小：小四equal pt12
        header_run.font.size = Pt(12)
        # 设置加粗：是
        header_run.font.bold = True
        # 设置字体颜色：黑色
        header_run.font.color.rgb = RGBColor(0, 0, 0)
        # 设置下划线： False
        header_run.font.underline = False


def set_body_font(pargraph):
    runs = pargraph.runs

    for header_run in runs:
        # 设置西文字体：Times New Roman
        header_run.font.name = 'Times New Roman'
        # 设置中文字体：宋体
        header_run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 设置字体大小：小四equal pt12
        header_run.font.size = Pt(12)
        # 设置加粗：是
        header_run.font.bold = False
        # 设置字体颜色：黑色
        header_run.font.color.rgb = RGBColor(0, 0, 0)
        # 设置下划线： False
        header_run.font.underline = False



class SummaryExtracter:
    def __init__(self, file=None, plate=None):
        self.file = file
        self.plate = plate
        self.data = None

    def __del__(self):
        pass

    def update_file_and_plate(self, file, plate):
        self.file = file
        self.plate = plate
        self.data = None

    def get_df(self):
        df = pd.read_csv(self.file, sep='\t')

        df_target = df[['Sample Filename', 'affymetrix-plate-peg-wellposition', 'Pass/Fail', 'DQC', 'QC call_rate',
                        'call_rate', 'QC computed_gender']]

        self.data = df_target.copy()
        self.data.sort_values(by='affymetrix-plate-peg-wellposition', inplace=True)

        self.data["DQC"] = round(df['DQC'], 3)
        self.data["QC call_rate"] = round(df['QC call_rate'], 3)
        self.data["call_rate"] = round(df['call_rate'], 3)

        return self.data

    def get_size(self):
        return len(self.data)

    def get_fail_list(self):
        fail_list = []

        col_lst = self.data['Pass/Fail'].tolist()
        for ind, val in enumerate(col_lst):
            if val == 'Fail':
                fail_list.append(ind+1)
        return fail_list

    def get_fail_counter(self):
        risk_dic = {}

        col_lst = self.data['Pass/Fail'].tolist()
        while np.nan in col_lst:
            col_lst.remove(np.nan)
        risk_dic = Counter(col_lst)
        return risk_dic

    def get_dqc_fail_list(self):
        dqc_fail_list = []

        col_lst = self.data['DQC'].tolist()
        for ind, val in enumerate(col_lst):
            if val < 0.82:
                dqc_fail_list.append(ind + 1)
        return dqc_fail_list

    def get_qc_fail_list(self):
        qc_fail_list = []

        col_lst = self.data['QC call_rate'].tolist()
        for ind, val in enumerate(col_lst):
            if val < 97:
                qc_fail_list.append(ind + 1)
        return qc_fail_list

    def get_avg_qc_callrate(self):
        df = self.data[['Pass/Fail', 'QC call_rate']]

        df = df[df['Pass/Fail'] == 'Pass']
        # print(df)

        val = df['QC call_rate'].mean().round(3)
        return val

def set_paragraph_font(pargraph):
    runs = pargraph.runs

    for header_run in runs:
        # 设置字体格式
        header_run.font.name = 'Times New Roman'  # 注：这个好像设置 run 中的西文字体
        # 设置中文字体 需导入 qn 模块
        header_run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 设置字体大小
        header_run.font.size = Pt(12)   #小四equal pt12
        # 设置加粗
        header_run.font.bold = True
        # 设置字体颜色 需导入 rgb 颜色模块
        header_run.font.color.rgb = RGBColor(0, 0, 0)
        # 设置下划线
        # header_run.font.underline = True


def set_body_font(pargraph):
    runs = pargraph.runs

    for header_run in runs:
        # 设置西文字体：Times New Roman
        header_run.font.name = 'Times New Roman'
        # 设置中文字体：宋体
        header_run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 设置字体大小：小四equal pt12
        header_run.font.size = Pt(12)
        # 设置加粗：是
        header_run.font.bold = False
        # 设置字体颜色：黑色
        header_run.font.color.rgb = RGBColor(0, 0, 0)
        # 设置下划线： False
        header_run.font.underline = False


def init_table(document, cols):
    tabl = document.add_table(rows=1, cols=0, style='Table Grid')
    tabl.autofit = False
    # 设置列宽
    for ite in cols:
        tabl.add_column(width=Mm(cols[ite]))
    # 设置表头
    hdr_cells = tabl.rows[0].cells
    for i, ite in enumerate(cols):
        hdr_cells[i].text = ite
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    return tabl


def pad_table(table, df_data):
    #t填充数据
    for index, row in df_data.iterrows():
        row_cells = table.add_row().cells  # add new row
        i = 0
        for item in row:
            item_str = str(item)
            if item_str == "nan":
                item_str = ''
            row_cells[i].text = item_str
            i += 1

    # 设置字体
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(9)


if __name__ == '__main__':
    file_name: str = 'C:/Users/ben/PycharmProjects/TZ0028/0028.txt'

    global data

    try:
        summary = SummaryExtracter()
        summary.update_file_and_plate(file_name, 'TZ0028')

        data = summary.get_df()
        print(data)

        dic = summary.get_fail_counter()
        print(dic)

        num = summary.get_size()
        print(num)

        ls = summary.get_fail_list()
        print(ls)

        ls = summary.get_dqc_fail_list()
        print(ls)

        ls = summary.get_qc_fail_list()
        print(ls)

        val = summary.get_avg_qc_callrate()
        print(val)

    except Exception as e:
        print(str(e))

    if 1:
        doc = Document()

        # # 增加表格
        # summary_cols = {"Sample Filename": 30,
        #             "affymetrix-plate-peg-wellposition": 30,
        #             " Pass/Fail": 22,
        #             "DQC": 20,
        #             "QC call_rate": 15,
        #             "call_rate": 20,
        #             "QC computed_gender": 20}
        #
        # summary_table = init_table(doc, summary_cols)
        #
        # pad_table(summary_table, data)

        # 增加Title：
        title_str = 'CAS-CN1基因芯片实验报告TZ0028'
        title = doc.add_heading(title_str, 0)
        title.style = doc.styles['Normal']
        set_title_font(title)

        title.paragraph_format.space_before = Pt(0)
        title.paragraph_format.space_after = Pt(0)
        title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        # 段落正文
        time_par = doc.add_paragraph('报告时间：')
        set_heading_font(time_par)
        time_str = time.strftime("%Y%m%d", time.localtime())
        time_run = time_par.add_run(time_str)
        # 日期加下划线
        time_run.font.bold = False
        time_run.font.underline = True


        # 增加标题1：
        heading = doc.add_heading('一、实验基本信息', 1)
        set_heading_font(heading)
        # 增加标题2：
        heading = doc.add_heading('1.1 样本类型', 2)
        set_heading_font(heading)
        # 段落正文
        body = doc.add_paragraph('基因组DNA样本。')
        set_body_font(body)
        # 增加标题2：
        heading = doc.add_heading('1.2 芯片类型', 2)
        set_heading_font(heading)
        body = doc.add_paragraph('CAS-CN1 96Array Plate。')
        # 增加标题2：
        heading = heading = doc.add_heading('1.3 实验流程', 2)
        set_heading_font(heading)


        # 保存文件
        doc.save(f'CAS-CN1基因芯片实验报告-testing1.docx')



