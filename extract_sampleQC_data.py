import pandas as pd
from collections import Counter
import numpy as np
from docx import Document
from docx.shared import Cm, Mm


class SampleQcExtracter:
    def __init__(self, file=None, plate=None):
        self.file = file
        self.plate = plate
        self.data = None

    def __del__(self):
        pass

    def update_file_and_plate(self, file, plate):
        self.file = file
        self.plate = plate
        self.data = None

    def get_df(self):
        df = pd.read_excel(io=self.file)

        self.data = df[df['新孔板号'] == self.plate]
        # self.data = df[df['新孔板号'] == self.plate].reset_index()

        self.data = self.data.copy()

        self.data['Picogreen浓度(ng/ul)'] = round(self.data['Picogreen浓度(ng/ul)'], 2)
        self.data['Nano浓度(ng/ul)'] = round(self.data['Nano浓度(ng/ul)'], 2)
        self.data['OD260/OD280'] = round(self.data['OD260/OD280'], 2)
        self.data['OD260/OD230'] = round(self.data['OD260/OD230'], 2)

        # self.data.round({'Picogreen浓度(ng/ul)': 2, 'Nano浓度(ng/ul)': 2, 'OD260/OD280': 2, 'OD260/OD230': 2})

        return self.data

    def get_size(self):
        return len(self.data)

    def get_risk_list(self):
        risk_list = []

        col_lst = self.data['结果'].tolist()
        for ind, val in enumerate(col_lst):
            if val == '风险上机':
                risk_list.append(ind+1)
        return risk_list

    def get_risk_counter(self):
        risk_dic = {}

        col_lst = self.data['备注'].tolist()
        while np.nan in col_lst:
            col_lst.remove(np.nan)
        risk_dic = Counter(col_lst)

        return risk_dic


if __name__ == '__main__':
    file_name = 'C:/Users/ben/PycharmProjects/样本质控汇总.xlsx'

    try:
        sample_qc = SampleQcExtracter()

        sample_qc.update_file_and_plate(file_name, 'TZ0028')

        sample_df = sample_qc.get_df()
        sample_num = sample_qc.get_size()
        sample_risk_list = sample_qc.get_risk_list()
        sample_risk_num = len(sample_risk_list)
        sample_risk_dic = sample_qc.get_risk_counter()

    except Exception as e:
        print(str(e))

    if 1:
        doc = Document()  # 新建文档对象
        # 增加表格
        colnames = ["孔板号", "孔号", "样本编号", "Picogreen浓度(ng/ul)", "Nano浓度(ng/ul)", "OD260/OD280",
                    "OD260/OD230", "备注", "结果"]
        colwidth = [18, 12, 25, 18, 15, 13, 13, 40, 20]
        # tablestyle在模板里指定
        sample_table = doc.add_table(rows=1, cols=0, style='Table Grid')
        sample_table.autofit = False
        # 设置列宽
        for i in range(len(colnames)):
            sample_table.add_column(width=Mm(colwidth[i]))

        # 设置表头
        hdr_cells = sample_table.rows[0].cells
        for i in range(len(colnames)):
            hdr_cells[i].text = colnames[i]

        # 填充数据
        for index, row in sample_df.iterrows():
            row_cells = sample_table.add_row().cells  # add new row
            i = 0
            for item in row:
                item_str = str(item)
                if item_str == "nan":
                    item_str = ''
                row_cells[i].text = item_str
                i += 1

        # 增加标题2：
        doc.add_heading('2.3质控结果', 2)

        sample_qc_comment = '样本浓度达标， OD260/OD280符合要求；'
        for key in sample_risk_dic:
            s = str(sample_risk_dic[key]) + '个样本' + str(key) + ';'
            sample_qc_comment += s
        s = f'风险上机，其余{sample_num - sample_risk_num}个样本均符合要求。'
        sample_qc_comment += s
        doc.add_paragraph(sample_qc_comment)


        # 保存文件
        doc.save(f'CAS-CN1基因芯片实验报告-testing.docx')

