from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

from docx.shared import RGBColor
from docx.shared import Pt, Cm, Mm
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn


import time

from PyQt5.QtCore import QThread
from PyQt5.QtCore import pyqtSignal

from extract_sampleQC_data import SampleQcExtracter
from extract_midQC_data import MidQcExtracter
from extract_summary_data import SummaryExtracter
from extract_plates_name import PlatesExtracter


def set_title_font(pargraph):
    pargraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    runs = pargraph.runs

    pargraph.paragraph_format.space_before = Pt(7.5)
    pargraph.paragraph_format.space_after = Pt(7.5)
    pargraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    for header_run in runs:
        # 设置西文字体：Times New Roman
        header_run.font.name = 'Times New Roman'
        # 设置中文字体：宋体
        header_run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 设置字体大小：小三equal pt15
        header_run.font.size = Pt(15)
        # 设置加粗：是
        header_run.font.bold = True
        # 设置字体颜色：黑色
        header_run.font.color.rgb = RGBColor(0, 0, 0)
        # 设置下划线： False
        header_run.font.underline = False


def set_heading_font(pargraph, color='black', size=12, bold=True):
    runs = pargraph.runs

    pargraph.paragraph_format.space_before = Pt(0)
    pargraph.paragraph_format.space_after = Pt(0)
    pargraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    for header_run in runs:
        # 设置西文字体：Times New Roman
        header_run.font.name = 'Times New Roman'
        # 设置中文字体：宋体
        header_run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 设置字体大小：小四equal pt12
        header_run.font.size = Pt(size)
        # 设置加粗：是
        header_run.font.bold = bold
        # 设置字体颜色：黑色
        if color == 'red':
            header_run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            header_run.font.color.rgb = RGBColor(0, 0, 0)
        # 设置下划线： False
        header_run.font.underline = False


def set_body_font(pargraph):
    runs = pargraph.runs

    pargraph.paragraph_format.space_before = Pt(0)
    pargraph.paragraph_format.space_after = Pt(0)
    pargraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    for header_run in runs:
        # 设置西文字体：Times New Roman
        header_run.font.name = 'Times New Roman'
        # 设置中文字体：宋体
        header_run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 设置字体大小：小四equal pt12
        header_run.font.size = Pt(12)
        # 设置加粗：是
        header_run.font.bold = False
        # 设置字体颜色：黑色
        header_run.font.color.rgb = RGBColor(0, 0, 0)
        # 设置下划线： False
        header_run.font.underline = False


def set_note_font(pargraph):
    runs = pargraph.runs

    pargraph.paragraph_format.space_before = Pt(0)
    pargraph.paragraph_format.space_after = Pt(0)
    pargraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    for header_run in runs:
        # 设置西文字体：Times New Roman
        header_run.font.name = 'Times New Roman'
        # 设置中文字体：宋体
        header_run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 设置字体大小：小四equal pt12
        header_run.font.size = Pt(9)
        # 设置加粗：是
        header_run.font.bold = False
        # 设置字体颜色：黑色
        header_run.font.color.rgb = RGBColor(0, 0, 0)
        # 设置下划线： False
        header_run.font.underline = False


def insert_paragraph_line(par, pos):
    p = par._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
                                  'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
                                  'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
                                  'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
                                  'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
                                  'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
                                  'w:pPrChange'
                                  )
    line = OxmlElement(f'w:{pos}')
    line.set(qn('w:val'), 'single')
    line.set(qn('w:sz'), '6')
    line.set(qn('w:space'), '1')
    line.set(qn('w:color'), 'auto')
    pBdr.append(line)


def init_table(document, cols, height_cm):
    tabl = document.add_table(rows=1, cols=0, style='Table Grid')
    tabl.autofit = False
    # 设置列宽
    for ite in cols:
        tabl.add_column(width=Mm(cols[ite]))
    # 设置表头
    hdr_cells = tabl.rows[0].cells
    for i, ite in enumerate(cols):
        hdr_cells[i].text = ite
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    hdr_row = tabl.rows[0]
    hdr_row.height = Cm(height_cm)
    return tabl


def pad_table(table, df_data, height_cm):
    # 填充数据
    for index, row in df_data.iterrows():
        row_line = table.add_row()
        row_line.height = Cm(height_cm)
        row_cells = row_line.cells  # add new row
        i = 0
        for item in row:
            item_str = str(item)
            if item_str == "nan":
                item_str = ''
            row_cells[i].text = item_str
            i += 1

    # 设置字体
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    font = run.font
                    # 设置字体格式
                    font.name = 'Times New Roman'
                    # 设置中文字体 需导入 qn 模块
                    font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    font.size = Pt(9)


def pad_table_add(table, df_data):
    # 填充数据
    for index, row in df_data.iterrows():
        row_line = table.add_row()
        row_line.height = Cm(1.0)
        row_cells = row_line.cells  # add new row
        i = 0
        for item in row:
            item_str = str(item)
            if item_str == "nan":
                item_str = ''
            row_cells[i].text = item_str
            i += 1

    # 设置字体
    for ind, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    font = run.font
                    # 设置西文字体
                    font.name = 'Times New Roman'
                    # 设置中文字体
                    font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    # 字体大小
                    font.size = Pt(9)
                    if ind == 5 or ind == 6:
                        # 设置字体颜色：黑色
                        font.color.rgb = RGBColor(255, 0, 0)

def set_table_bgcolor(table, col_cnt, col, color):
    shading_list = locals()
    for i in range(col_cnt):
        shading_list['shading_elm_' + str(i)] = \
            parse_xml(r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=color))
        table.rows[col].cells[i]._tc.get_or_add_tcPr().append(shading_list['shading_elm_' + str(i)])


class ReportGenerater(QThread):
    color_str = 'F2DBDB'
    gene_qc_req = ('1）双链定量浓度≥10ng/ul，总量≥600ng；',
                   '2）OD260/OD280 介于1.8至2.0之间，OD260/OD230≥1.5；',
                   '3）1%琼脂糖电泳90%DNA片段大小大于10Kb。')
    gene_qc_explain = ('1）合格：完全符合质量要求。',
                       '2）风险上机：5ng/ul≤浓度＜10ng/ul，300 ng≤总量＜600ng，DNA轻微降解，蛋白或RNA轻微污染。',
                       '3）不合格：浓度＜5ng/ul，总量＜300ng，DNA严重降解，蛋白或RNA污染严重。')
    experimental_process = ('1) 核对样本信息，样本入库。',
                            '2) 挑样，样本DNA分别进行Picogreen精确定量，Nano检测DNA纯度，琼脂糖电泳检测完整性。',
                            '3) 根据质控结果再次排板。',
                            '4) 样本DNA样本均一化，DNA扩增。',
                            '5) 扩增产物片段化、沉淀。',
                            '6) DNA干燥、重悬，质控。',
                            '7) 杂交、清洗、染色、扫描，数据分析。')
    sample_cols = {"孔板号": 14.5,
                   "孔号": 10.6,
                   "样本编号": 22,
                   "Picogreen浓度(ng/ul)": 15.9,
                   "Nano浓度(ng/ul)": 16.9,
                   "OD260/OD280": 15.1,
                   "OD260/OD230": 14.8,
                   "备注": 32,
                   "结果": 18.0}
    mid_cols = {"孔板号": 18.8,
                "孔号": 25.4,
                "样本编号": 28.3,
                "稀释后浓度（吸光度法）ng/μl": 41.3,
                "结果": 36.5}
    summary_cols = {"Sample Filename": 32,
                    "affymetrix-plate-peg-wellposition": 18.3,
                    " Pass/Fail": 18,
                    "DQC": 16.9,
                    "QC call_rate": 18,
                    "call_rate": 17.2,
                    "QC computed_gender": 28}
    plates_cols = {"起始板": 28.6,
                   "转移板": 28.6,
                   "孔板命名": 41.3,
                   "转板目的": 56.7, }
    log_signal = pyqtSignal(str)
    state_signal = pyqtSignal(str)

    def __init__(self, plate=None, sample_file=None, mid_file=None, summary_file=None, plates_file=None, pics_dir=None, parent=None):
        super(ReportGenerater, self).__init__(parent)

        self.__plate_name = plate
        self.__sampleQC_file_name = sample_file
        self.__midQC_file_name = mid_file
        self.__summary_file_name = summary_file
        self.__plates_file_name = plates_file
        self.__pics_dir = pics_dir

        self.sample_extracter = SampleQcExtracter()
        self.mid_extracter = MidQcExtracter()
        self.summary = SummaryExtracter()
        self.plates = PlatesExtracter()

    def __del__(self):
        self.wait()

    def update_file(self, plate, sample_file, mid_file, summary_file, plates_file, pics_dir):
        """更新各个文件路径。"""
        self.__plate_name = plate
        self.__sampleQC_file_name = sample_file
        self.__midQC_file_name = mid_file
        self.__summary_file_name = summary_file
        self.__plates_file_name = plates_file
        self.__pics_dir = pics_dir

    # def doit(self):
    def run(self):
        """线程中提取文件数据并生成报告。"""
        try:
            self.log_signal.emit("Extract data start. ")

            self.log_signal.emit(f"Extract sample QC data from {self.__sampleQC_file_name}.")
            self.sample_extracter.update_file_and_plate(self.__sampleQC_file_name, self.__plate_name)
            sample_df = self.sample_extracter.get_df()
            sample_num = self.sample_extracter.get_size()
            sample_risk_list = self.sample_extracter.get_risk_list()
            sample_risk_num = len(sample_risk_list)
            sample_risk_dic = self.sample_extracter.get_risk_counter()
            self.log_signal.emit("Extract sample QC data successfully.")

            self.log_signal.emit(f"Extract middle QC data from {self.__midQC_file_name}.")
            self.mid_extracter.update_file_and_plate(self.__midQC_file_name, self.__plate_name)
            mid_df = self.mid_extracter.get_df()
            mid_num = self.mid_extracter.get_size()
            mid_risk_list = self.mid_extracter.get_risk_list()
            mid_risk_num = len(mid_risk_list)
            mid_risk_dic = self.mid_extracter.get_risk_counter()
            self.log_signal.emit("Extract middle QC data successfully.")

            self.log_signal.emit(f"Extract summary data from {self.__summary_file_name}.")
            self.summary.update_file_and_plate(self.__summary_file_name, self.__plate_name)
            summary_df = self.summary.get_df()
            summary_fail_list = self.summary.get_fail_list()
            summary_dqc_fail_list = self.summary.get_dqc_fail_list()
            summary_qc_fail_list = self.summary.get_qc_fail_list()
            summary_average_qc_callrate = self.summary.get_avg_qc_callrate()
            self.log_signal.emit("Extract summary data successfully.")

            self.log_signal.emit(f"Extract plates data from {self.__plates_file_name}.")
            self.plates.update_file_and_plate(self.__plates_file_name, self.__plate_name)
            plates_df = self.plates.get_df()
            
            self.log_signal.emit("Extract plates data successfully.")

            self.log_signal.emit("Generate report start.")

            # 新建文档对象
            # doc = Document('C:/Users/ben/PycharmProjects/TZ0028/ReportTemplet.docx')
            doc = Document()

            sec = doc.sections
            sec0 = sec[0]
            header = sec0.header
            footer = sec0.footer
            header.is_linked_to_previous = False

            self.log_signal.emit('Generate document header.')
            # 设置页眉
            header0_par = header.paragraphs[0]
            insert_paragraph_line(header0_par, 'bottom')
            # 设置段落文本对齐
            header0_par.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # 页眉内容
            header_run = header0_par.add_run(
                '泰州国科医学检验实验室                                              暨泰州综保区保税研发实验室')
            # 设置字体格式
            header_run.font.name = 'Times New Roman'
            # 设置中文字体 需导入 qn 模块
            header_run.font.element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')
            # 设置字体大小
            header_run.font.size = Pt(10.5)
            # 设置加粗
            header_run.font.bold = True
            # 设置字体颜色 需导入 rgb 颜色模块
            header_run.font.color.rgb = RGBColor(47, 84, 150)
            # 设置下划线
            # header_run.font.underline = True

            self.log_signal.emit('Generate document footer.')
            # 设置页脚
            footer_par = footer.paragraphs[0]
            insert_paragraph_line(footer_par, 'top')
            # 设置段落文本对齐
            footer_par.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # 页脚内容
            font_run = footer_par.add_run(
                '公司地址：泰州市综合保税区标准厂房二期2号A座（A2）                       联系方式：0523-86669879')
            # 设置字体格式
            font_run.font.name = 'Times New Roman'
            # 设置中文字体 需导入 qn 模块
            # from docx.oxml.ns import qn
            font_run.font.element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')
            # 设置字体大小
            font_run.font.size = Pt(9)
            # 设置加粗
            font_run.font.bold = True
            # 设置字体颜色 需导入 rgb 颜色模块
            font_run.font.color.rgb = RGBColor(0, 0, 0)
            # 设置下划线
            # header_run.font.underline = True

            self.log_signal.emit('Generate document title.')
            # 增加Title：
            title_str = 'CAS-CN1基因芯片实验报告' + self.__plate_name
            title = doc.add_heading(title_str, 0)
            title.style = doc.styles['Normal']
            set_title_font(title)

            # 段落正文
            time_par = doc.add_paragraph('报告时间：')
            set_heading_font(time_par)
            time_str = time.strftime("%Y%m%d", time.localtime())
            time_run = time_par.add_run(time_str)
            # 日期加下划线
            time_run.font.bold = False
            time_run.font.underline = True

            self.log_signal.emit('Generate session1.')
            # 增加标题1：
            heading = doc.add_heading('一、实验基本信息', 1)
            set_heading_font(heading)
            # 增加标题2：
            heading = doc.add_heading('1.1 样本类型', 2)
            set_heading_font(heading)
            # 段落正文
            body = doc.add_paragraph('基因组DNA样本。')
            set_body_font(body)
            # 增加标题2：
            heading = doc.add_heading('1.2 芯片类型', 2)
            set_heading_font(heading)
            body = doc.add_paragraph('CAS-CN1 96Array Plate。')
            # 增加标题2：
            heading = heading = doc.add_heading('1.3 实验流程', 2)
            set_heading_font(heading)
            for ite in ReportGenerater.experimental_process:
                body = doc.add_paragraph(ite)
                set_body_font(body)

            self.log_signal.emit('Generate session2.')
            # 增加标题1：
            heading = doc.add_heading('二、样本DNA质控结果', 1)
            set_heading_font(heading)
            # 增加标题2：
            heading = doc.add_heading('2.1样本基因组DNA浓度汇总表', 2)
            set_heading_font(heading)
            # 增加表格
            sample_table = init_table(doc, ReportGenerater.sample_cols, 0.6)
            pad_table(sample_table, sample_df, 0.6)
            sample_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # 添加背景色
            self.log_signal.emit("Sample QC table, setting background color.")
            for col in sample_risk_list:
                set_table_bgcolor(sample_table, 9, col, ReportGenerater.color_str)

            # 增加标题2：
            heading = doc.add_heading('2.2 琼脂糖电泳检测图', 2)
            set_heading_font(heading)
            body = doc.add_paragraph('样本DNA质控胶图已提供，' + self.__plate_name + '是在已质控的样本中中挑选出的合格样本重新排版进行上机实验。')
            set_body_font(body)
            # 增加标题2：
            heading = doc.add_heading('2.3质控结果', 2)
            set_heading_font(heading)

            sample_qc_comment = ''
            if sample_risk_num == 0:
                sample_qc_comment = f'{sample_num}个样本均符合要求。'
            else:
                for key in sample_risk_dic:
                    s = str(sample_risk_dic[key]) + '个样本' + str(key) + ';'
                    sample_qc_comment += s
                s = f'风险上机，其余{sample_num - sample_risk_num}个样本均符合要求。'
                sample_qc_comment += s
            body = doc.add_paragraph(sample_qc_comment)
            set_body_font(body)

            # 增加标题2：
            heading = doc.add_heading('2.4 质控说明', 2)
            set_heading_font(heading)
            # 增加标题3：
            heading = doc.add_heading('1. 基因组DNA质量要求：', 3)
            set_heading_font(heading)
            for ite in ReportGenerater.gene_qc_req:
                body = doc.add_paragraph(ite)
                set_body_font(body)
            # 增加标题3：
            heading = doc.add_heading('2. 基因组DNA质检结果说明：', 3)
            set_heading_font(heading)
            for ite in ReportGenerater.gene_qc_explain:
                body = doc.add_paragraph(ite)
                set_body_font(body)

            self.log_signal.emit('Generate session3.')
            # 增加标题1：
            heading = doc.add_heading('三、杂交前片段化产物质控', 1)
            set_heading_font(heading)

            # 增加标题2：
            heading = doc.add_heading('3.1紫外吸光法检测浓度', 2)
            set_heading_font(heading)

            # # 增加表格
            mid_table = init_table(doc, ReportGenerater.mid_cols, 0.6)
            pad_table(mid_table, mid_df, 0.53)
            mid_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # 设置背景色
            for col in mid_risk_list:
                set_table_bgcolor(mid_table, 5, col, ReportGenerater.color_str)

            # 增加标题2：
            heading = doc.add_heading('3.2  3%琼脂糖凝胶电泳：', 2)
            set_heading_font(heading)

            # 增加图片
            pic_para = doc.add_paragraph()
            pic_run = pic_para.add_run('      ')
            pic_run.add_picture(self.__pics_dir+'/1.png', width=Cm(2.7), height=Cm(3.7))
            pic_para.add_run('  ')
            pic_run = pic_para.add_run()
            pic_run.add_picture(self.__pics_dir+'/2.png', width=Cm(4.5), height=Cm(3.7))
            pic_para.add_run('  ')
            pic_run = pic_para.add_run()
            pic_run.add_picture(self.__pics_dir+'/3.png', width=Cm(4.5), height=Cm(3.7))
            # 增加标题2：
            heading = doc.add_heading('3.3 质控结果', 2)
            set_heading_font(heading)
            par = doc.add_paragraph('100倍稀释后,')
            if len(mid_risk_list) == 0:
                par.add_run(f'{mid_num}个样本浓度均不低于80 ng/uL，符合要求。')
            else:
                par.add_run(f'{mid_risk_num}个样本浓度低于80 ng/uL，风险上机，其余{mid_num-mid_risk_num}个均合格。')

            # 增加标题2：
            heading = doc.add_heading('3.4 质控说明：', 2)
            set_heading_font(heading)
            body = doc.add_paragraph('1）100倍稀释后片段化DNA产物的浓度应＞80ng/uL；')
            set_body_font(body)
            body = doc.add_paragraph('2）片段化产物跑3%琼脂糖凝胶电泳结果主带明显，且分布于25-125bp。')
            set_body_font(body)

            #
            doc.add_page_break()

            self.log_signal.emit('Generate session4.')
            # 增加标题1：
            heading = doc.add_heading('四、数据检出率汇总', 1)
            set_heading_font(heading)
            # 增加标题2：
            heading = doc.add_heading('4.1 数据分析结果总结（AxiomAnalysisSuite软件导出）', 2)
            set_heading_font(heading)
            # 增加图片
            doc.add_picture(self.__pics_dir+'/4.png', width=Cm(15.4), height=Cm(11.4))
            doc.add_picture(self.__pics_dir+'/5.png', width=Cm(9.9), height=Cm(8.2))

            #
            doc.add_page_break()

            # 增加标题2：
            heading = doc.add_heading('4.2 数据分析结果汇总（AxiomAnalysisSuite软件导出）', 2)
            set_heading_font(heading)
            # 添加表格
            summary_table = init_table(doc, ReportGenerater.summary_cols, 0.6)
            pad_table(summary_table, summary_df, 0.48)
            summary_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for col in summary_fail_list:
                set_table_bgcolor(summary_table, 7, col, ReportGenerater.color_str)
            self.log_signal.emit('Generate session5.')
            # 增加标题1：
            heading = doc.add_heading('五、结果分析', 1)
            set_heading_font(heading)
            par = doc.add_paragraph('1.  ' + sample_qc_comment)
            set_body_font(par)

            par = doc.add_paragraph('2.  杂交前质控结果：100倍稀释后,')
            if len(mid_risk_list) == 0:
                par.add_run(f'{mid_num}个样本浓度均不低于80 ng/uL，符合要求。')
            else:
                par.add_run(f'{mid_risk_num}个样本浓度低于80 ng/uL，风险上机，其余{mid_num - mid_risk_num}个均合格。')
            par.add_run('片段大小分布满足要求。')
            set_body_font(par)

            par = doc.add_paragraph('3.  下机结果：')
            if len(summary_dqc_fail_list) == 0:
                par.add_run('整板样本DQC通过;')
            else:
                par.add_run(f'{len(summary_dqc_fail_list)}个样本DQC未通过;')
            if len(summary_qc_fail_list) == 0:
                par.add_run('整板样本QC通过;')
            else:
                par.add_run(f'{len(summary_qc_fail_list)}个样本QC未通过;')
            if len(summary_fail_list) == 0:
                par.add_run('整板检测结果样本通过率100%;')
            else:
                rate = (sample_num - len(summary_fail_list)) / sample_num * 100
                par.add_run(f'整板检测结果样本通过率{rate}%;')
            par.add_run(f'通过样本的平均QC call_rate值{summary_average_qc_callrate}%。')
            set_body_font(par)

            # 插入分页
            doc.add_page_break()

            # 增加标题1：
            heading = doc.add_heading('附录：', 1)
            set_heading_font(heading)
            heading = doc.add_paragraph('转板拍照记录')
            set_heading_font(heading)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # 增加表格
            plates_table = init_table(doc, ReportGenerater.plates_cols, 0.8)
            pad_table_add(plates_table, plates_df)
            plates_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            heading = doc.add_paragraph('注：红色标记为质控过程转板')
            set_note_font(heading)

            self.log_signal.emit('Generate session Addition.')
            # 增加标题2：
            heading = doc.add_heading('附图：', 2)
            set_heading_font(heading, color='Black', size=9)

            # 增加标题3：
            heading = doc.add_heading('1.蓝色冻存管到PCR板', 3)
            set_heading_font(heading, color='red', bold=False)
            doc.add_picture(self.__pics_dir+'/6.jpg', width=Cm(14.6), height=Cm(4.5))

            # 增加标题3：
            heading = doc.add_heading('2.PCR板到深孔板', 3)
            set_heading_font(heading, color='red', bold=False)
            doc.add_picture(self.__pics_dir+'/7.jpg', width=Cm(14.6), height=Cm(4.5))

            # 增加标题3：
            heading = doc.add_heading('3.深孔板到PCR', 3)
            set_heading_font(heading, color='red', bold=False)
            doc.add_picture(self.__pics_dir+'/8.jpg', width=Cm(14.6), height=Cm(4.5))

            # 增加标题3：
            heading = doc.add_heading('4.PCR板到酶标板', 3)
            set_heading_font(heading, color='red', bold=False)
            doc.add_picture(self.__pics_dir+'/9.jpg', width=Cm(14.6), height=Cm(4.5))

            # 增加标题3：
            heading = doc.add_heading('5.酶标板到点样板', 3)
            set_heading_font(heading, color='red', bold=False)
            doc.add_picture(self.__pics_dir+'/10.jpg', width=Cm(14.6), height=Cm(4.5))

            # 增加标题3：
            heading = doc.add_heading('6.PCR板到杂交板', 3)
            set_heading_font(heading, color='red', bold=False)
            doc.add_picture(self.__pics_dir+'/11.jpg', width=Cm(14.6), height=Cm(4.5))

            # 保存文件
            doc.save(f'CAS-CN1基因芯片实验报告{self.__plate_name}.docx')

            self.log_signal.emit("Save as "+f'CAS-CN1基因芯片实验报告{self.__plate_name}.docx')
            self.state_signal.emit("Generated report successfully.")
        except Exception as e:
            self.log_signal.emit(str(e))
            self.state_signal.emit("Failed.")


if __name__ == '__main__':

    file_name_sample = 'C:/Users/ben/PycharmProjects/样本质控汇总.xlsx'
    file_name_mid = 'C:/Users/ben/PycharmProjects/中间质控汇总 2021.3.25(1).xls'
    file_name_summary = 'C:/Users/ben/PycharmProjects/TZ0028/0028.txt'
    file_name_plates = 'C:/Users/ben/PycharmProjects/TZ0028/plates_name.txt'
    dir_path = 'C:/Users/ben/PycharmProjects/TZ0028'
    cur_plate_name = 'TZ0028'

    print('s0')
    rep = ReportGenerater()
    print('s1')
    rep.update_file(cur_plate_name, file_name_sample, file_name_mid, file_name_summary, file_name_plates, dir_path)

    print('s2')
    rep.doit()
    print('s3')
